from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Mm, Pt


ROOT = Path("/Users/amayranjan/Desktop/devsecops_project")
TEMPLATE = Path("/Users/amayranjan/Downloads/Sample Report for the 20Marks component.docx")
OUTPUT_DIR = ROOT / "report_output"
OUTPUT_PATH = OUTPUT_DIR / "RoomEase_CICD_Project_Report.docx"
SHOT_DIR = Path("/private/tmp/report_screenshots")


TEAM = [
    ("1.", "1MS23CS012", "Aditya Karan", "1ms23cs012@msrit.edu"),
    ("2.", "1MS23CS019", "Amay Ranjan", "1ms23cs019@msrit.edu"),
    ("3.", "1MS23CS054", "Deepanshu Gupta", "1ms23cs054@msrit.edu"),
]

TOC_PAGES = ["3", "5", "7", "8", "10", "12", "15", "16", "17", "18", "19"]


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, align=None, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def clear_paragraph(paragraph):
    paragraph.text = ""
    for run in paragraph.runs:
        run.text = ""


def add_page_break(anchor):
    p = anchor.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)
    format_paragraph(p)
    return p


def add_heading(anchor, text):
    p = anchor.insert_paragraph_before("")
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    format_paragraph(p, before=0, after=4)
    return p


def add_subheading(anchor, text):
    p = anchor.insert_paragraph_before("")
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    format_paragraph(p, before=2, after=2)
    return p


def add_body(anchor, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = anchor.insert_paragraph_before("")
    run = p.add_run(text)
    set_font(run, size=12)
    format_paragraph(p, after=2, line=1.0, align=align)
    return p


def add_bullet(anchor, text):
    p = anchor.insert_paragraph_before("")
    run = p.add_run(f"• {text}")
    set_font(run, size=12)
    format_paragraph(p, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    return p


def add_caption(anchor, text):
    p = anchor.insert_paragraph_before("")
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    format_paragraph(p, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_spacer(anchor, points=6):
    p = anchor.insert_paragraph_before("")
    format_paragraph(p, after=points)
    return p


def insert_picture_before(doc, anchor, image_path, width_inches):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    format_paragraph(p, after=2)
    anchor._p.addprevious(p._p)


def insert_table_before(doc, anchor, rows, cols, style="Table Grid"):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    anchor._p.addprevious(table._element)
    return table


def clone_row(table, row_idx):
    new_tr = deepcopy(table.rows[row_idx]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def set_table_cell_text(cell, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    format_paragraph(p, align=align)


def set_document_margins(doc):
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Inches(0.32)
        section.right_margin = Inches(0.76)
        section.top_margin = Inches(0.26)
        section.bottom_margin = Inches(0.19)
        pg_sz = section._sectPr.pgSz
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")


def fill_cover_and_tables(doc):
    doc.paragraphs[4].text = "CI/CD Pipeline Implementation                      Faculty Name: Brunda G"
    for run in doc.paragraphs[4].runs:
        set_font(run, size=12)

    team_table = doc.tables[0]
    for idx, data in enumerate(TEAM, start=2):
        if idx >= len(team_table.rows):
            clone_row(team_table, len(team_table.rows) - 1)
        for col, value in enumerate(data):
            align = WD_ALIGN_PARAGRAPH.LEFT if col in (2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            set_table_cell_text(team_table.rows[idx].cells[col], value, align=align)

    toc_table = doc.tables[1]
    for idx, page in enumerate(TOC_PAGES, start=1):
        set_table_cell_text(toc_table.rows[idx].cells[2], page)

    eval_table = doc.tables[2]
    while len(eval_table.rows) < 4:
        clone_row(eval_table, len(eval_table.rows) - 1)
    for idx, member in enumerate(TEAM, start=1):
        row = eval_table.rows[idx]
        set_table_cell_text(row.cells[0], member[1])
        set_table_cell_text(row.cells[1], member[2], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[2], "")
        set_table_cell_text(row.cells[3], "")
        set_table_cell_text(row.cells[4], "")

    doc.paragraphs[37].text = "Name: Brunda G"
    doc.paragraphs[38].text = "Designation: Assistant Professor"
    doc.paragraphs[39].text = "Department: CSE"
    doc.paragraphs[40].text = "Signature: "
    for idx in (37, 38, 39, 40):
        for run in doc.paragraphs[idx].runs:
            set_font(run, size=12)


def section_intro(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Introduction to CI/CD")
    add_body(
        anchor,
        "Continuous Integration and Continuous Deployment (CI/CD) is a core DevSecOps practice in which every meaningful code change is automatically validated, built, and prepared for release through a predictable workflow. In a traditional manual process, developers write code, run some checks locally, build deployment artifacts by hand, and then rely on ad hoc steps to publish updates. That approach is error-prone and difficult to audit. CI/CD replaces that uncertainty with automation so that each source-code change follows the same engineering path from commit to release candidate.",
    )
    add_body(
        anchor,
        "The selected project for this report is RoomEase, a hostel management system developed using Node.js, Express.js, MongoDB, Mongoose, HTML, CSS, and browser-side JavaScript. The project is designed to solve practical hostel-administration tasks such as student registration, room allocation visibility, entry and exit attendance, curfew enforcement, violation and fine tracking, and student complaint handling. In addition to application features, the repository also contains DevSecOps assets including a Dockerfile, a Docker Compose stack, GitHub Actions workflow automation, Kubernetes manifests, Prometheus configuration, and Grafana dashboard provisioning.",
    )
    add_body(
        anchor,
        "For RoomEase, CI/CD is especially important because the project contains more than a simple web server. It includes business logic, data models, environment-based configuration, container packaging, security middleware, health endpoints, observability components, and infrastructure definitions. Any change to one part of the system can affect the others. A disciplined pipeline ensures that these changes are validated in a controlled sequence before they reach users or administrators.",
    )
    add_body(
        anchor,
        "From a CI perspective, the objective is to continuously integrate small and frequent updates into a shared codebase while automatically checking correctness. From a CD perspective, the objective is to continuously produce a release-ready artifact, such as a Docker image, and then deliver it to a deployment platform in a repeatable manner. The DevSecOps dimension further extends this approach by embedding security and validation activities inside the same release flow rather than treating them as optional post-development tasks.",
    )
    add_body(
        anchor,
        "In the context of this report, RoomEase serves as a strong case study because it demonstrates how an academic software project can be transformed into a production-oriented engineering system. Instead of stopping at application development, the repository shows how modern teams secure code, automate testing, package workloads, validate manifests, and support scalable deployment and monitoring. Therefore, this report examines not only what the project does functionally, but also how it is prepared for reliable delivery.",
    )


def section_tools(doc, anchor):
    add_page_break(anchor)
    add_heading(anchor, "Selected Tools and Their Roles")
    add_body(anchor, "The RoomEase repository uses a combination of development, automation, deployment, and monitoring tools. Each tool addresses a particular engineering concern, and together they create a complete delivery ecosystem around the application.")
    add_body(anchor, "Rather than depending on a single platform for everything, the project follows a layered strategy. Source control and automation are handled through GitHub and GitHub Actions, artifact packaging is handled by Docker, image hosting is handled through GitHub Container Registry, deployment readiness is supported by Render and Kubernetes assets, and runtime visibility is provided through Prometheus and Grafana. This separation of concerns makes the overall system more maintainable and easier to reason about.")
    table = insert_table_before(doc, anchor, rows=1, cols=2)
    headers = [("Tool / Platform", True), ("Role in the Project", True)]
    for idx, (text, bold) in enumerate(headers):
        set_table_cell_text(table.rows[0].cells[idx], text, bold=bold)
    items = [
        ("GitHub", "Acts as the central source-control platform where code, workflow files, and infrastructure definitions are versioned and collaboratively maintained."),
        ("GitHub Actions", "Automates testing, auditing, image build and push, security scan, Kubernetes validation, and deployment triggering whenever changes reach the main integration path."),
        ("Node.js + Express.js", "Provide the backend runtime and application framework that power RoomEase APIs, static page hosting, middleware, and request handling."),
        ("MongoDB + Mongoose", "Persist application data such as students, attendance records, complaints, and violations using schema-backed document models."),
        ("Jest + Supertest", "Validate health endpoints, security headers, frontend routes, and metadata APIs before deployment so regressions are identified early."),
        ("Docker", "Packages the application into a repeatable runtime image using a multi-stage, lightweight, non-root build process."),
        ("GitHub Container Registry", "Stores tagged container images produced by the pipeline, making them available as versioned deployment artifacts."),
        ("Render", "Provides an easy production deployment target through an automated deploy-hook based release trigger."),
        ("Kubernetes", "Supplies production-ready orchestration manifests including Namespace, ConfigMap, Secret, Deployment, Service, Ingress, HPA, NetworkPolicy, and PodDisruptionBudget."),
        ("Prometheus + Grafana", "Collect, scrape, and visualize metrics exposed by the backend, improving observability and runtime diagnostics."),
        ("Docker Compose", "Creates a complete local multi-service environment containing the application, MongoDB, Prometheus, and Grafana."),
    ]
    for tool, role in items:
        row = table.add_row()
        set_table_cell_text(row.cells[0], tool, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[1], role, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(anchor, "The combination of these tools shows that RoomEase is not just a CRUD application. It has been organized as a delivery-capable system in which coding, validation, packaging, deployment, and observability are treated as connected responsibilities.")


def section_architecture(doc, anchor):
    add_page_break(anchor)
    add_heading(anchor, "Architecture of the CI/CD Pipeline")
    add_body(
        anchor,
        "The CI/CD architecture of RoomEase can be understood as a flow of trust-building stages. A developer begins by modifying application code, workflow files, frontend pages, or infrastructure manifests. Once these changes are committed and pushed to the repository, GitHub becomes the event source that activates the automation pipeline. This is the first transition point from local development into controlled shared validation.",
    )
    insert_picture_before(doc, anchor, SHOT_DIR / "architecture.png", 6.1)
    add_caption(anchor, "Figure 1: High-level CI/CD architecture derived from the RoomEase repository.")
    add_body(
        anchor,
        "The GitHub Actions workflow then serves as the central automation engine. It checks out the repository, restores the Node.js toolchain, installs dependencies, and executes the defined validation stages. These stages include testing, security auditing, image construction, vulnerability scanning, and Kubernetes manifest verification. Each stage increases confidence in the change set before it is promoted further.",
    )
    add_body(
        anchor,
        "After validation, the application is transformed into an immutable Docker artifact. This is a critical architectural step because deployment should operate on a known, versioned image rather than on raw source code. The generated image can be stored in a registry, referenced by deployment systems, and reused consistently across environments without introducing build drift.",
    )
    add_body(
        anchor,
        "On the delivery side, the project supports both a simpler production path and a more scalable orchestration path. Render provides an immediate web-service deployment route triggered from the pipeline. At the same time, the repository includes Kubernetes resources for teams that want stronger control over replicas, autoscaling, ingress, resource isolation, and rollout behavior. This dual-path design increases adaptability for different hosting contexts.",
    )
    add_body(
        anchor,
        "The architecture is completed by observability. Prometheus collects application metrics from the `/metrics` endpoint, while Grafana provides dashboards for visualization. Health and readiness endpoints make it possible for containers, platforms, and orchestration systems to determine whether the service is alive and whether it is ready to accept traffic. This means the architecture is not limited to delivery alone; it extends into post-deployment verification as well.",
    )
    add_body(
        anchor,
        "Overall, this architecture supports DevSecOps objectives because verification, packaging, deployment readiness, and observability are all defined as version-controlled assets. Every release candidate follows the same path, becomes easier to audit, and can be reproduced with much less operational uncertainty.",
    )


def section_installation(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Installation & Configuration Steps")
    add_body(anchor, "A successful CI/CD implementation depends on both tooling setup and correct environment configuration. In RoomEase, installation is not limited to downloading dependencies; it also involves preparing runtime secrets, local services, registry access, and deployment settings.")
    steps = [
        "Clone the repository and inspect the project structure so that application code, tests, workflow files, Docker configuration, Kubernetes manifests, and monitoring assets are clearly understood before execution.",
        "Install the Node.js dependencies using `npm ci` or `npm install`. The lockfile ensures deterministic dependency resolution, which is important for matching local behavior with automated CI runs.",
        "Configure environment variables such as `MONGO_URI`, `NODE_ENV`, `PORT`, `ADMIN_UID`, `ADMIN_PASS`, `RATE_LIMIT_WINDOW_MS`, and `RATE_LIMIT_MAX`. These values separate secrets and deployment-specific behavior from source code.",
        "For local development, start the multi-container stack defined in `docker-compose.yml`. This launches the application, MongoDB, Prometheus, and Grafana on a dedicated bridge network, allowing developers to test features and observability together.",
        "For CI/CD execution, ensure that GitHub repository permissions and secrets are configured correctly. The workflow depends on access for package publishing and may use `RENDER_DEPLOY_HOOK` to trigger automated deployment after successful build stages.",
        "When targeting Kubernetes, base64-encode the MongoDB connection string, update the image reference in the deployment manifest, and apply the namespace, configmap, secret, deployment, service, ingress, HPA, network policy, and PDB resources in the correct order.",
        "Verify the deployment using health endpoints, logs, pod status, and metrics so that setup is confirmed not only at installation time but also at service-readiness time.",
    ]
    for idx, step in enumerate(steps, start=1):
        add_body(anchor, f"{idx}. {step}")
    add_body(anchor, "These steps demonstrate an important DevSecOps principle: configuration is part of delivery. A pipeline becomes dependable only when application setup, infrastructure setup, and secret management are all handled deliberately.")


def section_implementation(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Implementation Details")
    add_subheading(anchor, "Application-layer implementation")
    add_body(
        anchor,
        "The backend implementation in `app.js` combines application logic with DevSecOps-focused middleware. The Express server initializes security and observability concerns at startup rather than treating them as later additions. Helmet is used to strengthen HTTP response headers, Morgan records incoming requests, rate limiting protects the server from abuse, and CORS plus JSON middleware support controlled client communication.",
    )
    add_body(
        anchor,
        "The backend also exposes meaningful operational endpoints. `/healthz` confirms that the service process is alive, while `/readyz` uses database readiness state to determine whether the application is actually prepared to handle production traffic. This distinction is important in modern deployment systems because an application may be running but still not be ready to serve requests safely.",
    )
    add_body(
        anchor,
        "Business functionality is implemented through models such as `Student`, `Attendance`, `Violation`, and `Complaint`. These models support domain-specific workflows such as resident registration, curfew rule enforcement, complaint creation, and violation-based fine calculation. Because the repository contains both business features and operational concerns in a single codebase, CI/CD becomes necessary to ensure changes do not accidentally break unrelated areas.",
    )
    add_body(
        anchor,
        "Prometheus instrumentation is another notable implementation detail. The project registers default metrics and custom metrics such as request duration, total HTTP requests, active requests, and database connection status. These metrics create a quantitative operational view of the system and can later be used for alerting, dashboarding, performance tuning, and troubleshooting.",
    )
    add_subheading(anchor, "Containerization and artifact implementation")
    add_body(
        anchor,
        "Containerization is defined in the `Dockerfile` using a deliberate two-stage build. In the dependency stage, the project installs production packages with `npm ci --omit=dev`, which improves repeatability and avoids unnecessary development-only modules in the final image. In the runtime stage, only the required application files and resolved dependencies are copied forward.",
    )
    add_body(
        anchor,
        "The runtime image uses Node 18 Alpine and runs the service as a non-root user. This is a strong security practice because it reduces the blast radius of a compromise inside the container. The Docker image also includes a native `HEALTHCHECK` that actively verifies the application through the `/healthz` endpoint, making it easier for runtime systems to identify unhealthy containers automatically.",
    )
    add_subheading(anchor, "Workflow and infrastructure implementation")
    add_body(
        anchor,
        "The CI/CD workflow stored in `.github/workflows/main.yml` is organized into five jobs. The first job prepares the Node.js environment, installs dependencies, runs an audit, and executes the test suite. This acts as the primary gatekeeper because there is no value in building or deploying an application that has already failed functional validation.",
    )
    add_body(
        anchor,
        "Once validation succeeds on the protected path, the workflow builds and pushes a Docker image to GitHub Container Registry. A separate security scan job then uses Trivy to examine the repository for high and critical vulnerabilities. In parallel, Kubernetes manifests are validated using `kubectl apply --dry-run=client`, which helps catch syntax or configuration issues before an actual cluster deployment attempt.",
    )
    add_body(
        anchor,
        "Operational readiness is further reinforced by the Docker Compose stack and the monitoring directory. Prometheus scrapes `/metrics`, Grafana provisions dashboards automatically, and the Kubernetes manifests add scaling, ingress routing, network restriction, and disruption control for more resilient deployment behavior. As a result, implementation is not isolated to the backend code alone; it spans the entire delivery lifecycle.",
    )


def section_screenshots(doc, anchor):
    add_page_break(anchor)
    add_heading(anchor, "Screenshots of Workflow and Output")
    add_body(anchor, "The following screenshots were prepared from the repository assets and local previews to demonstrate the implemented workflow, deployment flow, testing evidence, and application interface. Each screenshot supports a specific part of the report and shows how RoomEase combines functional features with DevSecOps practices.")
    insert_picture_before(doc, anchor, SHOT_DIR / "workflow-summary.png", 6.1)
    add_caption(anchor, "Figure 2: Summary of the five GitHub Actions jobs used in the RoomEase CI/CD pipeline.")
    add_body(anchor, "This figure summarizes the logical order of the CI/CD workflow. It shows that testing, image packaging, security analysis, manifest validation, and deployment triggering are all explicitly represented rather than being performed manually.")

    add_page_break(anchor)
    insert_picture_before(doc, anchor, SHOT_DIR / "admin-login.png", 5.3)
    add_caption(anchor, "Figure 3: RoomEase admin authentication screen used before entering the dashboard.")
    add_body(anchor, "The admin login interface is the entry point for secure administrative operations. It reflects the project requirement that sensitive management features should not be directly exposed without an authentication check.")
    insert_picture_before(doc, anchor, SHOT_DIR / "student-portal.png", 5.3)
    add_caption(anchor, "Figure 4: Student complaint portal provided by the RoomEase frontend.")
    add_body(anchor, "The student portal shows that the application is not only infrastructure-ready but also functionally complete from a user-facing perspective. Students can raise and track complaints using a dedicated interface that is separated from the admin dashboard.")

    add_page_break(anchor)
    insert_picture_before(doc, anchor, SHOT_DIR / "test-output.png", 5.8)
    add_caption(anchor, "Figure 5: Local Jest test result showing successful endpoint and route validation.")
    add_body(anchor, "The test-output screenshot acts as concrete validation evidence. It confirms that the project currently passes automated checks for health endpoints, security headers, served pages, and metadata output, which strengthens confidence in the pipeline.")
    insert_picture_before(doc, anchor, SHOT_DIR / "monitoring.png", 5.8)
    add_caption(anchor, "Figure 6: Docker Compose services used for application runtime, data, and monitoring.")
    add_body(anchor, "The monitoring summary reflects the local observability stack. It demonstrates that RoomEase is designed to run alongside infrastructure that measures health and performance, an important feature for any modern production-oriented application.")


def section_comparison(doc, anchor):
    add_page_break(anchor)
    add_heading(anchor, "Tool Comparisons (Optional)")
    add_body(anchor, "The project combines multiple tools because each one serves a different stage of software delivery. A comparison is useful because it explains not just what the project uses, but why those choices are reasonable for the given engineering goals.")
    table = insert_table_before(doc, anchor, rows=1, cols=4)
    for idx, text in enumerate(("Area", "Primary Tool", "Possible Alternative", "Why It Fits RoomEase")):
        set_table_cell_text(table.rows[0].cells[idx], text, bold=True)
    rows = [
        ("CI orchestration", "GitHub Actions", "Jenkins / GitLab CI", "GitHub Actions fits naturally because the repository is already hosted on GitHub and the workflow file is easy to maintain in source control."),
        ("Local runtime", "Docker Compose", "Manual service startup", "Docker Compose is better for this project because it launches the application, database, and monitoring tools together in one reproducible command."),
        ("Container registry", "GHCR", "Docker Hub", "GHCR integrates smoothly with GitHub Actions permissions and keeps source repository and image distribution within the same ecosystem."),
        ("Application deployment", "Render", "Railway / Fly.io / self-hosted VM", "Render gives a simpler deployment path for a student project while still supporting automated release triggering."),
        ("Scalable orchestration", "Kubernetes", "Docker Swarm / standalone containers", "Kubernetes is the stronger choice when autoscaling, service abstraction, ingress, and resilience policies are required."),
        ("Monitoring", "Prometheus + Grafana", "Hosted APM platforms", "This stack suits the project because it is open-source, widely adopted, and directly compatible with Prometheus metrics emitted by the application."),
    ]
    for area, tool, alt, reason in rows:
        row = table.add_row()
        set_table_cell_text(row.cells[0], area, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[1], tool, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[2], alt, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[3], reason, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(anchor, "These comparisons show that the RoomEase stack balances practicality and engineering maturity. The project uses approachable tools where simplicity is valuable and more advanced tools where scalability, security, or reliability justify the added complexity.")


def section_challenges(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Challenges Faced and Solutions")
    points = [
        "Database connectivity must differ between local development and production. The project solves this by reading `MONGO_URI` from environment variables rather than hardcoding it, which keeps secrets out of source control and allows different environments to use different data backends.",
        "A production-ready release requires more than just application code. RoomEase addresses this by storing Docker, monitoring, and Kubernetes assets in the same repository so the operational layer evolves together with the software layer.",
        "Security and reliability checks can be forgotten during manual delivery. GitHub Actions, dependency auditing, security scanning, and manifest validation automate these checks before deployment proceeds.",
        "Deployment targets often need health-based orchestration. The project solves this through `/healthz` and `/readyz` endpoints, Docker health checks, and readiness-aware orchestration support.",
        "Scaling and safe rollout are harder in a single-host environment. The Kubernetes manifests add readiness checks, autoscaling, ingress rules, network policy, and disruption budgeting to improve controlled operation under load or maintenance conditions.",
        "Observability is often an afterthought in student projects. RoomEase explicitly addresses this challenge by exposing Prometheus metrics and provisioning Grafana dashboards, making the system easier to inspect and troubleshoot after deployment.",
    ]
    for point in points:
        add_bullet(anchor, point)
    add_body(anchor, "These challenges show that the difficult part of DevSecOps is not merely writing a workflow file. The real work lies in designing an application and repository structure that support automation, security, traceability, and operational confidence from the beginning.")


def section_use_case_one(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Use Case Demonstration")
    add_body(
        anchor,
        "Consider a realistic use case in which the development team updates the complaint-handling experience so that students can report hostel issues more clearly and administrators can process them more efficiently. This kind of change might affect frontend forms, backend request handling, schema validation, and even deployment artifacts if dependencies or routes are modified.",
    )
    add_body(
        anchor,
        "After the updated code is pushed to the main branch, the GitHub Actions workflow begins automatically. The pipeline first checks out the repository, restores the Node.js environment, installs dependencies, and runs the automated test suite. This early gate is crucial because it prevents the team from spending time packaging and distributing changes that already violate expected behavior.",
    )
    add_body(
        anchor,
        "If validation succeeds, the workflow packages the updated RoomEase application into a Docker image and publishes it to GitHub Container Registry. This guarantees that the same release artifact can be consumed by deployment targets without rebuilding the application differently in each environment. In other words, the artifact tested by the pipeline is the same artifact later delivered to runtime infrastructure.",
    )


def section_use_case_two(anchor):
    add_page_break(anchor)
    add_heading(anchor, "Use Case Demonstration")
    add_body(
        anchor,
        "In the next stage of the same use case, the workflow continues beyond packaging and into release readiness. Trivy scans the repository for high and critical vulnerabilities, and Kubernetes manifests are validated through client-side dry runs. These actions ensure that the feature change is not only functionally acceptable but also operationally safe enough to promote.",
    )
    add_body(
        anchor,
        "Once the workflow reaches deployment, the Render deploy hook can trigger the production release path, while the Kubernetes manifests remain available for teams that want a cluster-based rollout. After deployment, administrators continue using the secured login interface to manage hostel operations, and students continue accessing the complaint portal shown earlier in the screenshots. The feature enhancement therefore reaches real users through an audited, automated, and repeatable release route.",
    )
    add_body(
        anchor,
        "Operationally, the application exposes `/healthz`, `/readyz`, and `/metrics`, allowing the hosting platform, Prometheus, and dashboard tooling to confirm availability after release. If a problem appears, logs, metrics, and health endpoints provide multiple signals that can guide diagnosis. This closes the DevSecOps loop by linking code change, validation, deployment, and runtime visibility into one continuous engineering lifecycle.",
    )


def section_references(anchor):
    add_page_break(anchor)
    add_heading(anchor, "References")
    refs = [
        "`PROJECT_EXPLAINED.md` - project overview, architecture explanation, and technology summary.",
        "`DEVOPS.md` - repository-level documentation for Docker, Kubernetes, CI/CD, security, and monitoring.",
        "`app.js` - backend application, middleware chain, health endpoints, metrics, and API logic.",
        "`package.json` - Node.js scripts, dependencies, and execution commands used in development and CI.",
        "`.github/workflows/main.yml` - GitHub Actions CI/CD pipeline definition.",
        "`Dockerfile` - multi-stage production image build configuration with non-root runtime execution.",
        "`docker-compose.yml` - local development stack with MongoDB, Prometheus, and Grafana.",
        "`k8s/README.md` and `k8s/*.yaml` - Kubernetes deployment, scaling, networking, and validation assets.",
        "`monitoring/prometheus/prometheus.yml` and Grafana provisioning files - observability configuration.",
        "`tests/health.test.js` - automated validation of health, headers, routes, and metadata APIs.",
        "`public/landing.html`, `public/admin-login.html`, `public/index.html`, and `public/student.html` - user interface components used by administrators and students.",
    ]
    for ref in refs:
        add_body(anchor, ref, align=WD_ALIGN_PARAGRAPH.LEFT)


def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document(str(TEMPLATE))
    set_document_margins(doc)
    fill_cover_and_tables(doc)

    anchor = next(p for p in doc.paragraphs if p.text.strip() == "Department of Computer Science and Engineering")
    section_intro(anchor)
    section_tools(doc, anchor)
    section_architecture(doc, anchor)
    section_installation(anchor)
    section_implementation(anchor)
    section_screenshots(doc, anchor)
    section_comparison(doc, anchor)
    section_challenges(anchor)
    section_use_case_one(anchor)
    section_use_case_two(anchor)
    section_references(anchor)
    add_page_break(anchor)

    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
